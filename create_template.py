from docx import Document

doc = Document()
doc.add_heading('Portaria CCM', 0)

p = doc.add_paragraph('A escola ')
p.add_run('{{ESCOLA}}').bold = True
p.add_run(' foi selecionada para o programa.')

doc.save('Modelo.docx')
print("Modelo.docx criado com sucesso.")
