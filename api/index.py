from flask import Flask, request, send_file, jsonify, render_template
import os
import re
import io
import json
from collections import defaultdict

# Imports seguros para prevenir falhas no arranque
try:
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    print(f"Startup Warning: {e}")

app = Flask(__name__)

# Configuração de Credenciais
SCOPES = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
SERVICE_ACCOUNT_FILE = 'credentials.json'

def get_services():
    creds = None
    if os.path.exists(SERVICE_ACCOUNT_FILE):
        creds = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    elif os.getenv('GOOGLE_CREDENTIALS'):
        creds_info = json.loads(os.getenv('GOOGLE_CREDENTIALS'))
        creds = service_account.Credentials.from_service_account_info(
            creds_info, scopes=SCOPES)
            
    if not creds:
        raise Exception("Credenciais não encontradas! Configura GOOGLE_CREDENTIALS ou credentials.json.")

    sheets_service = build('sheets', 'v4', credentials=creds)
    drive_service = build('drive', 'v3', credentials=creds)
    return sheets_service, drive_service

def extrair_id(link):
    m = re.search(r"/d/([a-zA-Z0-9-]+)", link)
    return m.group(1) if m else link.strip()

@app.route('/')
def home():
    try:
        return render_template('index.html')
    except Exception as e:
        return f"Erro ao renderizar template: {str(e)}", 500

@app.route('/api/ler-colunas', methods=['POST'])
def ler_colunas():
    try:
        data = request.json
        link = data.get('link')
        aba = data.get('aba')
        sheets_service, _ = get_services()
        sid = extrair_id(link)
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=sid, range=f"'{aba}'!A4:ZZ4").execute()
        valores = result.get('values', [[]])
        if not valores or not valores[0]:
            return {"error": "Linha 4 vazia ou aba não encontrada"}, 400
        cabecalho = valores[0]
        colunas = [f"{i}|{n.strip().replace('\"', '')}" for i, n in enumerate(cabecalho) if n.strip()]
        return {"colunas": colunas}
    except Exception as e:
        return {"error": str(e)}, 500

@app.route('/api/processar', methods=['POST'])
def processar():
    try:
        data = request.json
        link = data.get('link')
        aba = data.get('aba')
        letra_escola = data.get('letra_escola')
        filtro_excluir = data.get('filtro_excluir')
        colunas_remover_str = data.get('colunas_remover', '')
        formato = data.get('formato')

        sheets_service, drive_service = get_services()
        sid = extrair_id(link)

        # Lendo todos os dados
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=sid, range=f"'{aba}'!A:ZZ").execute()
        linhas_todas = result.get('values', [])

        if len(linhas_todas) < 4:
            return {"error": "Planilha curta demais (menos de 4 linhas)"}, 400

        cabecalho_original = linhas_todas[3]

        # --- ATUALIZAÇÃO: LÓGICA DE FILTRO GLOBAL ---
        # Divide os termos por vírgula para permitir múltiplos filtros (ex: FALSE, #REF!)
        termos_proibidos = [t.strip().upper() for t in filtro_excluir.split(',')] if filtro_excluir else []

        idx_esc = ord(letra_escola.upper()) - ord('A')
        idx_remover = [int(x) for x in colunas_remover_str.split(',') if x.strip().isdigit()]
        colunas_que_ficam = [(i, c.strip()) for i, c in enumerate(cabecalho_original) if i not in idx_remover and c.strip()]

        grupos = defaultdict(list)

        for linha in linhas_todas[4:]:
            if not linha or not any(str(c).strip() for c in linha): continue

            # Transforma a linha completa numa string para busca global
            linha_texto_completa = " ".join([str(celula).upper() for celula in linha])
            
            # Se encontrar qualquer termo proibido em QUALQUER coluna, pula a linha
            if any(termo in linha_texto_completa for termo in termos_proibidos if termo):
                continue

            # Verifica se as colunas mínimas de dados existem na linha
            if len(linha) < 2: continue

            dados_linha = []
            for i, _ in colunas_que_ficam:
                val = str(linha[i]).strip() if i < len(linha) else ""
                dados_linha.append(val)

            escola = str(linha[idx_esc]).strip() if idx_esc < len(linha) else "GERAL"
            grupos[escola].append(dados_linha)

        if not grupos:
             return {"error": f"❌ Nenhum dado encontrado. O filtro '{filtro_excluir}' eliminou tudo ou a aba está vazia."}, 400

        # --- GERAÇÃO WORD (MANTIDA) ---
        doc = Document()
        for section in doc.sections:
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

        # Cabeçalho do Documento
        h_txt = ("SECRETARIA DE ESTADO DE EDUCAÇÃO\nSECRETARIA ADJUNTA DE GESTÃO DE PESSOAS\n"
                 "DIRETORIA DE ORGANIZAÇÃO DE PESSOAL\nCOORDENADORIA DE CONTROLE E MOVIMENTAÇÃO")
        p_h = doc.add_paragraph()
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_h = p_h.add_run(h_txt)
        run_h.bold = True
        run_h.font.size = Pt(11)

        doc.add_paragraph().add_run("\nPORTARIA N° 000006-2026 - SAGEP").bold = True
        
        # Texto Padrão
        corpo = ("\nA Secretária Adjunta... RESOLVE:\n\nArt. 1º Ficam concedidas Férias...\n\n"
                 "Art. 2º Esta Portaria entra em vigor na data de publicação.\n\nBelém (PA), 2026.")
        doc.add_paragraph(corpo).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

        # Tabelas por Escola
        for escola, lista_dados in grupos.items():
            doc.add_paragraph().add_run(f"ANEXO - {escola}").bold = True
            tabela = doc.add_table(rows=1, cols=len(colunas_que_ficam))
            tabela.style = 'Table Grid'
            
            for j, (_, nome) in enumerate(colunas_que_ficam):
                tabela.rows[0].cells[j].text = nome
            
            for reg in lista_dados:
                row_cells = tabela.add_row().cells
                for j, val in enumerate(reg):
                    row_cells[j].text = str(val)

        # Salvar e Enviar
        base_dir = "/tmp" if not os.name == 'nt' else os.getcwd()
        path_docx = os.path.join(base_dir, "Portaria_2026.docx")
        doc.save(path_docx)

        return send_file(path_docx, as_attachment=True, download_name="Portaria_2026.docx")

    except Exception as e:
        return {"error": str(e)}, 500

if __name__ == '__main__':
    app.run(debug=True)
